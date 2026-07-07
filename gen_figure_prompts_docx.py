"""生成四张说明书配图的中文提示词（Word）。"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT = Path(r"c:\Users\32872\Desktop\etf智能体\说明书配图生成提示词.docx")

FIGURES: list[dict[str, str]] = [
    {
        "no": "图3-1",
        "title": "ETF 交易池与动态进攻池",
        "manual_ref": "说明书 3.1 节 · 图3-1",
        "purpose": "展示 10 只固定稳健池 + 3 只条件启用进攻池，以及宽基趋势闸门逻辑。",
        "style": (
            "扁平信息图（infographic），白底或浅灰底，16:9 横版，"
            "适合插入 Word 说明书。配色：宽基蓝、商品金、红利绿、行业紫、进攻池橙红。"
            "无照片、无 3D、无卡通人物。所有文字清晰可读、简体中文。"
        ),
        "must_include": (
            "【固定稳健池 10 只，分四组】\n"
            "宽基（5）：510300 沪深300ETF、510050 上证50ETF、510500 中证500ETF、"
            "510330 华夏沪深300ETF、159338 中证A500ETF\n"
            "商品避险（2）：518880 黄金ETF、159985 豆粕ETF\n"
            "红利防御（1）：510880 红利ETF\n"
            "行业（2）：512880 证券ETF、512010 医药ETF\n\n"
            "【动态进攻池 3 只】（虚线框，标注「条件启用」）\n"
            "159915 创业板ETF、588000 科创50ETF、159949 创业板50ETF\n\n"
            "【启用条件闸门】（醒目箭头或开关图标）\n"
            "宽基复合趋势分均值 ≥ +3% 时并入候选池\n"
            "（参考：510300 + 159915 + 588000 的「5日涨 + 3日涨×0.5」均值）\n\n"
            "【底部说明】常态仅用固定池；强势行情临时扩展至 13 只候选，再经评分选 1～3 只"
        ),
        "prompt_zh": (
            "一张专业的中文金融科技信息图，标题「ETF 交易池与动态进攻池」。"
            "画面分为上下两层：上层是「固定稳健池」四个彩色分区卡片，分别列出 10 只 ETF 代码与简称；"
            "下层是「动态进攻池」三个 ETF，用橙色虚线边框包裹，标注「≥+3% 条件启用」。"
            "中间有一个醒目的闸门/开关元素，文字为「宽基复合趋势分 ≥ +3%」。"
            "箭头从闸门指向进攻池，表示条件满足时并入候选。"
            "底部一行小字：「合并候选池 → 综合评分排序 → 最终持有 1～3 只」。"
            "风格：扁平矢量、商务答辩 PPT 质感、白底、细线图标、分区色块清晰、"
            "所有中文标签大号加粗、无英文水印、无股票 K 线图照片。"
        ),
        "prompt_en": (
            "Professional Chinese fintech infographic, title "
            "\"ETF Trading Pool and Dynamic Offensive Pool\", 16:9 landscape, flat vector style, "
            "white background. Top section: four colored blocks for fixed pool "
            "(10 ETFs in categories: broad-base, commodities, dividend, sector). "
            "Bottom section: dashed orange box with 3 offensive ETFs, label "
            "\"enabled when composite trend score >= +3%\". "
            "Central gate/switch icon connecting layers. Clean business presentation aesthetic, "
            "all labels in simplified Chinese, highly readable typography, no photos, no 3D."
        ),
        "negative": "模糊文字、乱码中文、英文标题占主导、写实股票交易大厅、人物插画、水印、低分辨率。",
        "size": "建议 1920×1080 或 2560×1440，导出 PNG，300dpi 用于印刷可另存 TIFF。",
    },
    {
        "no": "图3-2",
        "title": "关键词词典与三级筛选",
        "manual_ref": "说明书 3.2.1 节 · 图3-2",
        "purpose": "说明新闻如何从原始标题正文，经拒稿、主题匹配、催化剂匹配到强/弱/拒绝三级结果。",
        "style": (
            "左到右流程图 + 右侧小型「主题词典样例」卡片。"
            "白底，主色蓝绿，拒绝步骤用红色，弱信号用黄色，强信号用绿色。16:9 横版。"
        ),
        "must_include": (
            "【流程五步】\n"
            "① 原始新闻标题/正文\n"
            "② 前置拒稿：资金榜、主力榜、融资榜、盘中异动、成交额榜 → 直接丢弃\n"
            "③ 赛道匹配：ETF 主题词典（新闻命中某 ETF 主题关键词）\n"
            "④ 催化剂匹配：六类实质催化剂\n"
            "   政策落地 | 资本开支 | 订单交付 | 业绩数据 | 技术突破 | 资金流向\n"
            "⑤ 信号分级：\n"
            "   强信号（主题+催化剂）基础分 0.35～0.42\n"
            "   弱信号（仅主题）基础分 0.16\n"
            "   拒绝（无主题或无实质内容）\n\n"
            "【右侧样例卡片，仅示例 3 只 ETF】\n"
            "512880：券商、证券、降准\n"
            "518880：黄金、避险、美联储\n"
            "512010：医药、创新药、医保"
        ),
        "prompt_zh": (
            "一张中文新闻筛选流程信息图，标题「关键词词典与三级筛选」。"
            "主画面为从左到右的五步水平流程图：原始新闻 → 前置拒稿（红色叉号）"
            "→ ETF 主题词典匹配 → 六类催化剂匹配 → 信号分级输出。"
            "第三步旁画一本打开的小词典图标；第四步用六个小标签展示催化剂类别；"
            "第五步分出三条支路：绿色「强信号」、黄色「弱信号」、红色「拒绝」。"
            "画面右侧竖排三个小卡片，标题「主题词典样例」，列出三只 ETF 的关键词示例。"
            "扁平商务风格，箭头清晰，所有文字为简体中文，适合学术论文插图，无人物无照片。"
        ),
        "prompt_en": (
            "Chinese infographic titled \"Keyword Dictionary and Three-Level Screening\", "
            "horizontal left-to-right flowchart on white background. Five steps: raw news, "
            "pre-filter rejection (red X), ETF theme keyword matching, six catalyst categories, "
            "signal grading (strong/weak/reject in green/yellow/red). "
            "Right sidebar: three sample ETF keyword cards. Flat vector, academic report style, "
            "simplified Chinese labels only, crisp arrows, no characters or photos."
        ),
        "negative": "流程顺序颠倒、催化剂类别缺失、全英文标签、花哨渐变背景、文字过小无法辨认。",
        "size": "建议 1920×1080，横向；若 Word 单栏排版可用 4:3 竖版另存一版。",
    },
    {
        "no": "图3-3",
        "title": "新闻信号加工全流程",
        "manual_ref": "说明书 3.2 节（新鲜/陈旧切分 + 两级筛选）",
        "purpose": "展示 15:00 切分双池、关键词预筛、DeepSeek 语义评分及汇入综合评分。",
        "style": (
            "双通道流水线图，中间有时间轴「昨日 15:00」切分标记。"
            "蓝=新鲜池，灰=陈旧池，底部强调降噪比例。16:9 横版，扁平矢量。"
        ),
        "must_include": (
            "【顶部】数百条原始财经新闻（可写「每日 500+ 条」示意）\n"
            "【切分线】昨日 15:00 收盘 → 左侧「新鲜新闻池」、右侧「陈旧新闻池」\n"
            "【每条通道相同步骤】\n"
            "关键词规则预筛（压缩至 ≤20 条）→ DeepSeek 语义评分（每批 ≤15 条）\n"
            "→ 产出主题原始分 → 映射为 0～100 主题分\n"
            "【汇流】新鲜主题分（权重 25%）+ 陈旧主题分（权重 10%）→ 汇入「综合评分」\n"
            "【底部标注】降噪比例超过 95%（数百条 → ≤20 条进入大模型）"
        ),
        "prompt_zh": (
            "一张中文双通道数据处理流水线信息图，标题「新闻信号加工全流程」。"
            "顶部宽条表示「原始财经新闻（500+ 条）」，中间有一条醒目的竖向时间分界线，"
            "标注「昨日 15:00 切分」。分界线左侧蓝色通道「新鲜新闻池」，右侧灰色通道「陈旧新闻池」。"
            "每条通道向下经过三个模块：关键词预筛、DeepSeek 语义评分、主题分输出。"
            "两条通道在底部汇合到一个方框「综合评分」，旁注「新鲜 25% + 陈旧 10%」。"
            "最底部横幅文字「降噪 >95%：数百条 → ≤20 条进入大模型」。"
            "风格简洁专业，适合技术说明书，简体中文，无照片无 3D。"
        ),
        "prompt_en": (
            "Chinese dual-pipeline data processing infographic, title "
            "\"News Signal Processing Pipeline\", 16:9, flat vector. "
            "Top: raw finance news bar. Center: vertical split at \"Yesterday 15:00\". "
            "Left blue path: fresh news pool; right gray path: stale news pool. "
            "Each path: keyword pre-filter, DeepSeek semantic scoring, theme score output. "
            "Merge at bottom into \"Composite Score\" with weights 25% and 10%. "
            "Footer banner: \"Noise reduction >95%\". Simplified Chinese, clean academic style."
        ),
        "negative": "单通道无分叉、缺少 15:00 切分、权重数字错误、混乱配色、英文为主。",
        "size": "建议 1920×1080；答辩 PPT 可直接全屏展示。",
    },
    {
        "no": "图3-4",
        "title": "多层递进式仓位风控",
        "manual_ref": "说明书 3.7.3 节 · 表3-9",
        "purpose": "将宽基评估、新闻调仓、经济日历、评分闸门、单 ETF 上限五层风控可视化为自上而下漏斗。",
        "style": (
            "自上而下漏斗/叠层图，每层一个色带，越往下约束越紧。"
            "白底，金融风控主题，可用盾牌/刹车隐喻但保持扁平。16:9 或 3:4 竖版均可。"
        ),
        "must_include": (
            "【第1层 宽基市场评估】（最宽）\n"
            "沪深300+创业板+科创50 复合趋势分 → 仓位比例\n"
            "≤−5% 或 ≤−2%：15% 试探仓；≤−0.5%：40%；中性约 70%；"
            "0.5%～2%：85%；≥2%：90%\n\n"
            "【第2层 新闻情绪调仓】\n"
            "按置信度/催化剂调整；低置信（<0.19）额外 ×0.62\n"
            "持仓只数：无新闻或低强度 1 只；中等 2 只；强信号最多 3 只\n\n"
            "【第3层 经济日历分级】\n"
            "高影响 1～2 条 → 上限 85%；3～5 条 → 75%；6+ 条 → 65%\n"
            "日历总条数为 0 → 总仓位硬顶 50%\n\n"
            "【第4层 评分闸门】\n"
            "默认 50 分；大模型强信号（|态度分|≥0.5）可降至 42 分；不达标 → 空仓\n\n"
            "【第5层 单 ETF 集中度】（最窄）\n"
            "任一 ETF 不超过总资金 30%\n\n"
            "【侧注】大模型输出 stay_cash → 直接空仓，绕过买入"
        ),
        "prompt_zh": (
            "一张中文多层风控漏斗信息图，标题「多层递进式仓位风控」。"
            "画面为自上而下逐渐变窄的五层梯形漏斗，每层一条色带并附简短中文说明。"
            "第1层最宽「宽基市场评估」列出主要档位与仓位百分比；"
            "第2层「新闻情绪调仓」含持仓只数与低置信降仓；"
            "第3层「经济日历分级」85/75/65 与空日历 50% 硬顶；"
            "第4层「评分闸门」50 分/42 分；"
            "第5层最窄「单 ETF ≤30%」。"
            "左侧小图标链条表示「层层叠加、互不替代」。"
            "可选侧边红色旁路标注「大模型 stay_cash → 直接空仓」。"
            "扁平商务风，简体中文，清晰数字，适合答辩，无写实交易所场景。"
        ),
        "prompt_en": (
            "Chinese multi-layer risk control funnel infographic, title "
            "\"Layered Position Risk Controls\", top-down narrowing five tiers. "
            "Tier1 widest: broad-market regime caps; tier2: news sentiment adjustment; "
            "tier3: economic calendar tiers 85/75/65 and 50% hard cap; "
            "tier4: score gate 50/42; tier5 narrowest: single ETF 30% max. "
            "Side note: LLM stay_cash forces empty position. Flat vector, simplified Chinese, "
            "readable percentages, professional fintech report style."
        ),
        "negative": "写「≤−4% 空仓」、层数不足五层、档位数字与上文不符、3D 金属盾牌占满画面、文字模糊。",
        "size": "建议 1920×1080 横版；若插图占半页可用 1200×1600 竖版漏斗。",
    },
]


def set_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)


def add_label_para(doc: Document, label: str, text: str, *, bold_label: bool = True) -> None:
    p = doc.add_paragraph()
    r = p.add_run(f"{label}：")
    r.bold = bold_label
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r2 = p.add_run(text)
    r2.font.name = "宋体"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_codeblock(doc: Document, text: str) -> None:
    for line in text.split("\n"):
        p = doc.add_paragraph(line if line else " ")
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def main() -> Path:
    doc = Document()
    set_font(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ETF 智能体说明书配图 · 生成提示词")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    doc.add_paragraph(
        "用途：将下列提示词复制到文生图工具（如 Midjourney、DALL·E、即梦、可灵等）"
        "或交给设计师，生成说明书四张结构示意图。\n"
        "对照基准：etf_agent 当前生产代码与 0612(2).docx 修订口径。"
    )

    doc.add_heading("通用设置（四张图共用）", level=2)
    add_label_para(
        doc,
        "画面比例",
        "优先 16:9 横版（1920×1080）；插入 Word 单栏时可按需裁切。",
    )
    add_label_para(
        doc,
        "整体风格",
        "扁平矢量信息图 / 商务答辩 PPT 插图；白底或极浅灰底；禁止照片写实、3D 渲染、卡通人物。",
    )
    add_label_para(
        doc,
        "文字要求",
        "所有界面文字使用简体中文；ETF 代码、百分比、阈值必须与「必须包含的文字」一致。",
    )
    add_label_para(
        doc,
        "生成建议",
        "先用「中文主提示词」生成；不满意时用「英文备用提示词」；"
        "务必附带「负向提示词」避免乱码与风格跑偏。生成后检查数字再插入 Word。",
    )

    for fig in FIGURES:
        doc.add_page_break()
        h = doc.add_heading(f"{fig['no']}  {fig['title']}", level=1)
        for run in h.runs:
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

        add_label_para(doc, "说明书位置", fig["manual_ref"])
        add_label_para(doc, "配图目的", fig["purpose"])
        add_label_para(doc, "视觉风格", fig["style"])
        add_label_para(doc, "输出尺寸", fig["size"])

        doc.add_heading("必须包含的文字与数据（生成后逐项核对）", level=3)
        add_codeblock(doc, fig["must_include"])

        doc.add_heading("中文主提示词（推荐直接复制）", level=3)
        add_codeblock(doc, fig["prompt_zh"])

        doc.add_heading("英文备用提示词", level=3)
        add_codeblock(doc, fig["prompt_en"])

        doc.add_heading("负向提示词", level=3)
        add_codeblock(doc, fig["negative"])

    doc.add_page_break()
    doc.add_heading("插入说明书时的图题建议", level=2)
    rows = [
        ("图3-1", "ETF 交易池（10 只固定池 + 3 只动态进攻池）"),
        ("图3-2", "关键词词典与三级筛选流程"),
        ("图3-3", "新闻信号加工全流程（新鲜/陈旧双池）"),
        ("图3-4", "多层递进式仓位风控"),
    ]
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "图号"
    tbl.rows[0].cells[1].text = "图题"
    for i, (no, caption) in enumerate(rows, 1):
        tbl.rows[i].cells[0].text = no
        tbl.rows[i].cells[1].text = caption

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"已生成: {OUT}")
    return OUT


if __name__ == "__main__":
    main()
